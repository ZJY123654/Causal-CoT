from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

import pandas as pd


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "results" / "data preparation"
FIG = OUT / "figures"
TAB = OUT / "tables"
ORG = OUT / "origin_data"


def read_jsonl(path: Path):
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                yield json.loads(line)


def parse_count(text: str, patterns: list[str]) -> int:
    value = 0
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            value = max(value, int(match.group(1)))
    return value


def public_text(text: str) -> str:
    value = str(text or "")
    if not value:
        return value
    value = re.sub(
        r"[\u4e00-\u9fffA-Za-z0-9（）()·\-]{2,40}(?:有限责任公司|有限公司|集团有限公司|集团公司|工程局|施工局|项目经理部|项目部)",
        "施工单位A",
        value,
    )
    value = re.sub(r"(?:中国水利水电|中铁|葛洲坝|水电)[\u4e00-\u9fffA-Za-z0-9（）()·\-]{0,30}", "施工单位A", value)
    value = re.sub(r"[\u4e00-\u9fff]{2,8}(?:水电站|电站|水库|水站)", "某水电站", value)
    value = re.sub(r"[\u4e00-\u9fffA-Za-z0-9]{1,12}(?:标段|工程)", "某标段", value)
    value = re.sub(r"[\u4e00-\u9fff]{1,12}供电局", "相关单位A", value)
    value = re.sub(r"[\u4e00-\u9fff]{2,8}局", "相关单位A", value)
    value = re.sub(r"[\u4e00-\u9fff]{1,3}[一二三四五六七八九十\d]+路", "某线路", value)
    value = re.sub(
        r"(工人|民工|班长|副班长|队长|副队长|负责人|副主任|主任|指挥|司机|操作员|作业人员|施工人员)([\u4e00-\u9fff]{2,4})(?=[在因将把对从，。；、\\s])",
        r"\1A",
        value,
    )
    surname = "赵钱孙李周吴郑王冯陈褚卫蒋沈韩杨朱秦尤许何吕施张孔曹严华金魏陶姜谢邹喻柏水窦章云苏潘葛奚范彭郎鲁韦昌马苗凤花方俞任袁柳鲍史唐费廉岑薛雷贺倪汤滕殷罗毕郝邬安常乐于时傅皮卞齐康伍余元卜顾孟平黄和穆萧尹"
    value = re.sub(rf"(?<![\u4e00-\u9fff])([{surname}][\u4e00-\u9fff]{{1,2}})(?=[在因将把对从，。；、\\s])", "作业人员A", value)
    value = re.sub(rf"([{surname}])所站位置", "作业人员A所站位置", value)
    value = value.replace("违章指挥A", "违章指挥")
    value = value.replace("指挥A", "指挥")
    return value


def build_case_dataframe() -> pd.DataFrame:
    cases = list(read_jsonl(ROOT / "data" / "processed" / "anonymized_cases.jsonl"))
    rows = []
    for case in cases:
        raw_text = case.get("raw_text", "") or ""
        combined_text = " ".join(str(case.get(k, "")) for k in ["date", "date_time", "raw_text", "title"])
        year_match = re.search(r"(19\d{2}|20\d{2})", combined_text)
        consequence_text = " ".join(
            str(case.get(k, "")) for k in ["consequence_text", "economic_loss_text", "raw_text"]
        )
        death_count = parse_count(
            consequence_text,
            [r"死亡\s*(\d+)\s*人", r"(\d+)\s*人死亡", r"造成\s*(\d+)\s*人死亡"],
        )
        injury_count = parse_count(
            consequence_text,
            [r"重伤\s*(\d+)\s*人", r"(\d+)\s*人重伤", r"受伤\s*(\d+)\s*人", r"伤\s*(\d+)\s*人"],
        )
        rows.append(
            {
                "case_id": case.get("case_id", ""),
                "source_file": case.get("source_file", ""),
                "source_case_no": case.get("source_case_no", ""),
                "title": public_text(case.get("title", "")),
                "accident_type": case.get("accident_type") or "未标明",
                "year": int(year_match.group(1)) if year_match else "",
                "date_time": case.get("date_time") or case.get("date") or "",
                "location": public_text(case.get("location", "")),
                "death_count": death_count,
                "injury_count": injury_count,
                "text_chars": len(raw_text),
                "has_process": bool(case.get("process_text") or raw_text),
                "has_cause": bool(
                    case.get("cause_text") or case.get("direct_cause_text") or case.get("indirect_cause_text")
                ),
                "has_consequence": bool(case.get("consequence_text")),
                "has_measures": bool(case.get("measures_text") or "预防措施" in raw_text),
                "process_excerpt": public_text((case.get("process_text") or raw_text)[:120].replace("\n", " ")),
                "cause_excerpt": public_text(
                    (case.get("cause_text") or case.get("direct_cause_text") or "")[:120].replace("\n", " ")
                ),
                "measure_excerpt": public_text((case.get("measures_text") or "")[:120].replace("\n", " ")),
            }
        )
    return pd.DataFrame(rows)


def make_summary(df: pd.DataFrame, entities: list[dict], triples: list[dict], raw_entities: list[dict], raw_triples: list[dict]) -> dict:
    years = pd.to_numeric(df["year"], errors="coerce")
    summary = {
        "case_count": len(df),
        "source_file_count": int(df["source_file"].nunique()),
        "year_min": int(years.min()),
        "year_max": int(years.max()),
        "total_deaths_extracted": int(df["death_count"].sum()),
        "total_injuries_extracted": int(df["injury_count"].sum()),
        "mean_text_chars": round(float(df["text_chars"].mean()), 2),
        "median_text_chars": round(float(df["text_chars"].median()), 2),
        "raw_entity_count": len(raw_entities),
        "raw_triple_count": len(raw_triples),
        "canonical_entity_count": len(entities),
        "canonical_triple_count": len(triples),
    }
    for path in [ROOT / "data" / "fusion" / "fusion_report.json"]:
        if path.exists():
            summary.update(json.loads(path.read_text(encoding="utf-8")))
    anonymization_report = ROOT / "data" / "privacy" / "anonymization_report.json"
    if anonymization_report.exists():
        report = json.loads(anonymization_report.read_text(encoding="utf-8"))
        summary["anonymized_person_replacements"] = report["replacement_counts"]["person"]
        summary["anonymized_org_replacements"] = report["replacement_counts"]["organization"]
        summary["anonymized_project_replacements"] = report["replacement_counts"]["project"]
    extraction_progress = ROOT / "data" / "kg" / "extraction_progress.json"
    if extraction_progress.exists():
        progress = json.loads(extraction_progress.read_text(encoding="utf-8"))
        summary["llm_completed_cases"] = progress.get("completed_cases")
        summary["llm_total_elapsed_hms"] = progress.get("total_elapsed_hms")
        summary["llm_average_case_elapsed_hms"] = progress.get("average_case_elapsed_hms")
    return summary


def save_csvs(df: pd.DataFrame, entities: list[dict], triples: list[dict], summary: dict) -> dict[str, pd.DataFrame]:
    tables: dict[str, pd.DataFrame] = {}
    tables["dataset_summary"] = pd.DataFrame([{"metric": k, "value": v} for k, v in summary.items()])
    tables["accident_type_distribution"] = (
        df["accident_type"].fillna("未标明").replace("", "未标明").value_counts().reset_index()
    )
    tables["accident_type_distribution"].columns = ["accident_type", "case_count"]
    tables["source_distribution"] = df["source_file"].value_counts().reset_index()
    tables["source_distribution"].columns = ["source_file", "case_count"]
    year_df = df[df["year"] != ""].assign(year=lambda x: x["year"].astype(int))
    tables["year_distribution"] = year_df["year"].value_counts().sort_index().reset_index()
    tables["year_distribution"].columns = ["year", "case_count"]
    tables["canonical_entity_type_distribution"] = pd.DataFrame(
        Counter(entity.get("label", "Unknown") for entity in entities).most_common(),
        columns=["entity_type", "count"],
    )
    tables["canonical_relation_type_distribution"] = pd.DataFrame(
        Counter(triple.get("type", "Unknown") for triple in triples).most_common(),
        columns=["relation_type", "count"],
    )
    text_bins = pd.cut(df["text_chars"], bins=[0, 500, 1000, 1500, 2000, 3000, 5000, 100000], right=False)
    text_dist = df.groupby(text_bins, observed=False).size().reset_index(name="case_count")
    text_dist["text_length_bin"] = text_dist["text_chars"].astype(str)
    tables["text_length_distribution"] = text_dist[["text_length_bin", "case_count"]]
    tables["field_coverage"] = pd.DataFrame(
        [
            {"field": "accident_type", "coverage_count": int((df["accident_type"] != "未标明").sum())},
            {"field": "year/date", "coverage_count": int((df["year"] != "").sum())},
            {"field": "location", "coverage_count": int((df["location"].astype(str).str.len() > 0).sum())},
            {"field": "accident_process/raw_text", "coverage_count": int(df["has_process"].sum())},
            {"field": "cause_text", "coverage_count": int(df["has_cause"].sum())},
            {"field": "consequence_text", "coverage_count": int(df["has_consequence"].sum())},
            {"field": "preventive_measures", "coverage_count": int(df["has_measures"].sum())},
        ]
    )
    tables["field_coverage"]["coverage_percent"] = (tables["field_coverage"]["coverage_count"] / len(df) * 100).round(2)
    tables["retained_fields"] = pd.DataFrame(
        [
            ["case_id", "案例唯一编号", "链接清洗数据、LLM抽取结果、KG节点与因果矩阵"],
            ["source_file/source_case_no", "来源文件与原始案例编号", "数据追溯和复核"],
            ["title", "事故标题", "案例识别与 AccidentCase 节点名称"],
            ["accident_type", "事故类型", "类型统计、控制变量和 AccidentType 节点"],
            ["date/date_time", "事故时间", "年份分布与时序统计"],
            ["location", "事故地点/工程场景", "脱敏后保留泛化地点，用于场景分析"],
            ["process_text/raw_text", "事故经过", "证据片段切分和事故演化抽取"],
            ["cause_text/direct_cause_text/indirect_cause_text", "事故原因", "24Model致因实体和关系抽取"],
            ["consequence_text/economic_loss_text", "事故后果", "严重后果变量和后果节点"],
            ["measures_text", "整改与预防措施", "PreventiveMeasure 节点和 controlledBy 关系"],
            ["privacy_status", "脱敏状态", "标记公开数据是否完成不可逆脱敏"],
        ],
        columns=["field", "content", "purpose"],
    )
    tables["data_processing_flow"] = pd.DataFrame(
        [
            ["资料收集", "从安全管理网、应急管理部公开信息及《水电工程施工事故案例汇编》等来源汇集事故案例文本。", "原始 Word/网页文本"],
            ["格式转换", "读取 docx，并通过 Word COM 支持 doc 文档转换/读取。", "段落文本"],
            ["文本清洗", "统一空白、标点、页码和标题格式；按案例编号切分。", "cleaned_cases.jsonl"],
            ["字段解析", "解析标题、事故类型、时间、地点、伤亡、经过、原因、措施等字段。", "结构化案例记录"],
            ["不可逆脱敏", "替换人名、公司/项目部、施工局等敏感表达，不保存可逆映射。", "anonymized_cases.jsonl"],
            ["LLM抽取", "按24Model问题链执行证据、实体、关系、投票和验证。", "extraction_results.jsonl"],
            ["知识融合", "字符串召回、向量重排、NIL检测、同类型对齐、冲突记录和关系合并。", "canonical_entities/triples.jsonl"],
        ],
        columns=["step", "operation", "output"],
    )
    preferred = ["坍塌", "高处坠落", "机械伤害", "起重伤害", "触电", "物体打击", "车辆伤害", "淹溺", "其他伤害"]
    examples, used = [], set()
    for accident_type in preferred:
        candidates = df[df["accident_type"].astype(str).str.contains(accident_type, na=False)]
        if not candidates.empty:
            row = candidates.iloc[0].to_dict()
            examples.append(row)
            used.add(row["case_id"])
    for _, row in df[~df["case_id"].isin(used)].head(max(0, 8 - len(examples))).iterrows():
        examples.append(row.to_dict())
    tables["representative_cases"] = pd.DataFrame(examples[:8])[
        [
            "case_id",
            "source_case_no",
            "title",
            "accident_type",
            "date_time",
            "location",
            "death_count",
            "injury_count",
            "process_excerpt",
            "cause_excerpt",
            "measure_excerpt",
        ]
    ]
    for name, frame in tables.items():
        frame.to_csv(TAB / f"{name}.csv", index=False, encoding="utf-8-sig")
        if name not in {"retained_fields", "data_processing_flow", "representative_cases", "dataset_summary"}:
            frame.to_csv(ORG / f"{name}_for_origin.csv", index=False, encoding="utf-8-sig")
    return tables


def save_matplotlib_figures(tables: dict[str, pd.DataFrame]) -> None:
    import matplotlib.pyplot as plt

    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    def bar(data: pd.DataFrame, x: str, y: str, title: str, xlabel: str, ylabel: str, path: Path, top: int | None = None):
        d = data.head(top).copy() if top else data.copy()
        fig, ax = plt.subplots(figsize=(10, 6), dpi=160)
        ax.bar(d[x].astype(str), d[y], color="#4C78A8")
        ax.set_title(title)
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.tick_params(axis="x", rotation=30)
        for label in ax.get_xticklabels():
            label.set_ha("right")
        fig.tight_layout()
        fig.savefig(path, dpi=300)
        plt.close(fig)

    bar(tables["accident_type_distribution"], "accident_type", "case_count", "事故类型分布", "事故类型", "案例数", FIG / "fig_accident_type_distribution.png", 15)
    bar(tables["source_distribution"], "source_file", "case_count", "来源文件分布", "来源文件", "案例数", FIG / "fig_source_distribution.png")
    bar(tables["canonical_entity_type_distribution"], "entity_type", "count", "融合后实体类型分布", "实体类型", "实体数", FIG / "fig_entity_type_distribution.png", 20)
    bar(tables["canonical_relation_type_distribution"], "relation_type", "count", "融合后关系类型分布", "关系类型", "关系数", FIG / "fig_relation_type_distribution.png", 20)
    bar(tables["field_coverage"], "field", "coverage_percent", "字段覆盖率", "字段", "覆盖率(%)", FIG / "fig_field_coverage.png")
    if not tables["year_distribution"].empty:
        fig, ax = plt.subplots(figsize=(10, 5), dpi=160)
        ax.plot(tables["year_distribution"]["year"], tables["year_distribution"]["case_count"], marker="o", color="#F58518")
        ax.set_title("事故案例年份分布")
        ax.set_xlabel("年份")
        ax.set_ylabel("案例数")
        fig.tight_layout()
        fig.savefig(FIG / "fig_year_distribution.png", dpi=300)
        plt.close(fig)
    bar(tables["text_length_distribution"], "text_length_bin", "case_count", "案例文本长度分布", "字符数区间", "案例数", FIG / "fig_text_length_distribution.png")


def markdown_table(frame: pd.DataFrame, max_rows: int | None = None) -> str:
    data = frame.head(max_rows).copy() if max_rows else frame.copy()
    return data.to_markdown(index=False)


def save_manuscript(tables: dict[str, pd.DataFrame], summary: dict) -> None:
    summary_text = "\n".join([f"- {key}: {value}" for key, value in summary.items()])
    md = f"""# 4.1.1 Data preparation

## 数据来源与收集

本研究的数据集由三类资料整合形成：（1）安全管理网公开发布的安全生产资料和事故案例栏目；（2）中华人民共和国应急管理部公开发布的事故调查、事故通报和安全生产信息；（3）长江三峡技术经济发展有限公司组织编制的《水电工程施工事故案例汇编》。安全管理网主页列有“事故案例”等栏目，适合作为公开事故案例资料来源之一；应急管理部网站为国家应急管理与安全生产信息发布平台，可用于补充官方事故调查和通报信息。研究过程中将上述来源中与水利水电工程建设、施工活动、工程对象和事故致因有关的案例汇集为原始 Word 文档与结构化文本输入。

## 数据处理流程

{markdown_table(tables["data_processing_flow"])}

本研究采用不可逆脱敏策略。处理对象包括 `raw_text`、`title`、`location`、`cause_text`、`process_text`、`consequence_text`、`measures_text` 和证据字段等。真实人名被替换为“作业人员A/管理人员A/死者A/伤者A”等稳定占位词，公司、施工局、项目部、监理单位和建设单位等被替换为“施工单位A/监理单位A/建设单位A/项目部A”等泛化表达；工程地点保留为“某水电站”“某水库”“某标段”等非可逆泛化形式。脱敏过程不保存真实词与占位词之间的可逆映射。

## 最终数据规模

{summary_text}

最终得到 {summary["case_count"]} 条水利工程施工事故案例。经 LLM 抽取与知识融合后，原始图谱包含 {summary["raw_entity_count"]} 个实体和 {summary["raw_triple_count"]} 条关系，融合后得到 {summary["canonical_entity_count"]} 个 canonical 实体和 {summary["canonical_triple_count"]} 条 canonical 关系。LLM 抽取总耗时为 {summary.get("llm_total_elapsed_hms", "")}，平均每个案例耗时为 {summary.get("llm_average_case_elapsed_hms", "")}。

## 案例保留字段

{markdown_table(tables["retained_fields"])}

## 代表性案例示例

下表给出脱敏后的代表性案例。表中仅保留论文复核所需的摘要性字段；完整证据文本保存在 JSONL 文件中。

{markdown_table(tables["representative_cases"])}

## 数据统计分析

### 事故类型分布

{markdown_table(tables["accident_type_distribution"], 15)}

### 来源文件分布

{markdown_table(tables["source_distribution"])}

### 字段覆盖率

{markdown_table(tables["field_coverage"])}

### 融合后实体类型分布

{markdown_table(tables["canonical_entity_type_distribution"], 20)}

### 融合后关系类型分布

{markdown_table(tables["canonical_relation_type_distribution"], 20)}

## 图表文件

- `figures/fig_accident_type_distribution.png`
- `figures/fig_source_distribution.png`
- `figures/fig_year_distribution.png`
- `figures/fig_text_length_distribution.png`
- `figures/fig_field_coverage.png`
- `figures/fig_entity_type_distribution.png`
- `figures/fig_relation_type_distribution.png`

## Origin 数据文件

所有用于绘图的统计表均已另存为 `origin_data/*_for_origin.csv`，可直接导入 Origin 重新绘图和调整版式。
"""
    (OUT / "4.1.1_Data_preparation.md").write_text(md, encoding="utf-8")
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("4.1.1 Data preparation", level=1)
        doc.add_paragraph("本文件为 Markdown 正文的 Word 快速预览版；正式表格请使用 tables/ 下的 CSV。")
        for title, frame in [
            ("数据处理流程", tables["data_processing_flow"]),
            ("案例保留字段", tables["retained_fields"]),
            ("代表性案例示例", tables["representative_cases"]),
            ("事故类型分布", tables["accident_type_distribution"].head(15)),
            ("字段覆盖率", tables["field_coverage"]),
        ]:
            doc.add_heading(title, level=2)
            table = doc.add_table(rows=1, cols=len(frame.columns))
            for i, col in enumerate(frame.columns):
                table.rows[0].cells[i].text = str(col)
            for _, row in frame.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(frame.columns):
                    cells[i].text = str(row[col])[:500]
        doc.save(OUT / "4.1.1_Data_preparation.docx")
    except Exception as exc:
        (OUT / "docx_generation_error.txt").write_text(str(exc), encoding="utf-8")


def save_origin_project(tables: dict[str, pd.DataFrame]) -> dict:
    status = {"origin_project_saved": False, "message": ""}
    try:
        import originpro as op

        op.set_show(True)
        created_graphs = []
        for name, frame in tables.items():
            if name not in {
                "accident_type_distribution",
                "source_distribution",
                "year_distribution",
                "canonical_entity_type_distribution",
                "canonical_relation_type_distribution",
                "text_length_distribution",
                "field_coverage",
            }:
                continue
            workbook = op.new_book("w", lname=name)
            worksheet = workbook[0]
            worksheet.from_df(frame)
            if len(frame.columns) >= 2 and not frame.empty:
                graph_type = "line" if name == "year_distribution" else "col"
                try:
                    graph = op.new_graph(template="line" if graph_type == "line" else "column", lname=f"fig_{name}")
                    layer = graph[0]
                    layer.add_plot(worksheet, colx=0, coly=1, type=graph_type)
                    layer.rescale()
                    created_graphs.append(f"fig_{name}")
                except Exception as exc:
                    try:
                        plot_code = 200 if graph_type == "line" else 202
                        worksheet.lt_exec(f"worksheet -s 2 0 2 0; plotxy iy:=2 plot:={plot_code} ogl:=[<new name:=fig_{name}>];")
                        created_graphs.append(f"fig_{name}")
                    except Exception as inner_exc:
                        (OUT / f"origin_graph_error_{name}.txt").write_text(
                            f"Python API: {type(exc).__name__}: {exc}\nLabTalk: {type(inner_exc).__name__}: {inner_exc}",
                            encoding="utf-8",
                        )
        op.save(str(OUT / "data_preparation_origin_project.opju"))
        status["origin_project_saved"] = True
        status["message"] = f"Origin workbook and graph project saved. Created graphs: {', '.join(created_graphs)}"
    except Exception as exc:
        status["message"] = f"Origin automation unavailable: {type(exc).__name__}: {exc}"
        (OUT / "origin_automation_note.txt").write_text(status["message"], encoding="utf-8")
    return status


def main() -> None:
    for path in [OUT, FIG, TAB, ORG]:
        path.mkdir(parents=True, exist_ok=True)
    df = build_case_dataframe()
    entities = list(read_jsonl(ROOT / "data" / "fusion" / "canonical_entities.jsonl"))
    triples = list(read_jsonl(ROOT / "data" / "fusion" / "canonical_triples.jsonl"))
    raw_entities = list(read_jsonl(ROOT / "data" / "kg" / "entities.jsonl"))
    raw_triples = list(read_jsonl(ROOT / "data" / "kg" / "triples.jsonl"))
    summary = make_summary(df, entities, triples, raw_entities, raw_triples)
    tables = save_csvs(df, entities, triples, summary)
    save_matplotlib_figures(tables)
    save_manuscript(tables, summary)
    origin_status = save_origin_project(tables)
    manifest = {
        "output_dir": str(OUT),
        "files": [str(path.relative_to(OUT)) for path in OUT.rglob("*") if path.is_file()],
        "summary": summary,
        "origin_status": origin_status,
    }
    (OUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"output_dir": str(OUT), "file_count": len(manifest["files"]), "summary": summary, "origin_status": origin_status}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
