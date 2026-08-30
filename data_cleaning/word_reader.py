from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document

from .text_normalizer import normalize_paragraph


def read_docx_paragraphs(path: str | Path) -> list[str]:
    doc = Document(str(Path(path)))
    paragraphs = [normalize_paragraph(p.text) for p in doc.paragraphs if normalize_paragraph(p.text)]
    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_paragraph(cell.text) for cell in row.cells if normalize_paragraph(cell.text)]
            if cells:
                paragraphs.append(" | ".join(cells))
    return paragraphs


def read_doc_with_word_com(path: str | Path) -> list[str]:
    src = Path(path)
    with tempfile.TemporaryDirectory(prefix="hydraulic_doc_") as tmp:
        out_txt = Path(tmp) / "document.txt"
        ps = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{str(src).replace("'", "''")}')
$text = $doc.Content.Text
[System.IO.File]::WriteAllText('{str(out_txt).replace("'", "''")}', $text, [System.Text.Encoding]::UTF8)
$doc.Close($false)
$word.Quit()
"""
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(f"Word COM failed for {src}: {result.stderr.strip()}")
        text = out_txt.read_text(encoding="utf-8", errors="replace")
    return [normalize_paragraph(p) for p in text.splitlines() if normalize_paragraph(p)]


def read_word_paragraphs(path: str | Path) -> list[str]:
    src = Path(path)
    suffix = src.suffix.lower()
    if suffix == ".docx":
        return read_docx_paragraphs(src)
    if suffix == ".doc":
        return read_doc_with_word_com(src)
    raise ValueError(f"Unsupported Word format: {src}")
